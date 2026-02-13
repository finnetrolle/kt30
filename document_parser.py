"""
Document parser module for reading Microsoft Word and PDF files.
"""
import os
import logging
from typing import Optional
from docx import Document
from docx.opc.exceptions import PackageNotFoundError


logger = logging.getLogger(__name__)


class PDFParser:
    """Parser for PDF documents."""
    
    def __init__(self, file_path: str):
        """Initialize the PDF parser.
        
        Args:
            file_path: Path to the PDF document
        """
        self.file_path = file_path
        logger.debug(f"PDFParser initialized for: {file_path}")
    
    def extract_text(self) -> str:
        """Extract all text from the PDF document.
        
        Returns:
            Extracted text as a string
        """
        logger.info("Extracting text from PDF document...")
        
        try:
            # Try using pypdf first
            from pypdf import PdfReader
            
            reader = PdfReader(self.file_path)
            text_parts = []
            
            for i, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    text_parts.append(page_text)
                    logger.debug(f"  - Page {i+1}: {len(page_text)} characters")
            
            result = "\n\n".join(text_parts)
            logger.info(f"Total text extracted from PDF: {len(result)} characters")
            return result
            
        except ImportError:
            logger.warning("pypdf not available, trying PyPDF2...")
            
        try:
            # Fallback to PyPDF2
            from PyPDF2 import PdfReader
            
            reader = PdfReader(self.file_path)
            text_parts = []
            
            for i, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    text_parts.append(page_text)
                    logger.debug(f"  - Page {i+1}: {len(page_text)} characters")
            
            result = "\n\n".join(text_parts)
            logger.info(f"Total text extracted from PDF: {len(result)} characters")
            return result
            
        except ImportError:
            logger.error("No PDF library available. Please install pypdf or PyPDF2.")
            raise RuntimeError("No PDF library available. Please install pypdf or PyPDF2: pip install pypdf")
    
    def extract_structure(self) -> dict:
        """Extract document structure from PDF.
        
        Since PDFs don't have explicit heading structure like Word docs,
        we try to detect sections based on text patterns.
        
        Returns:
            Dictionary with document structure
        """
        logger.info("Extracting PDF document structure...")
        
        text = self.extract_text()
        lines = text.split('\n')
        
        structure = {
            "title": "",
            "sections": []
        }
        
        current_section = None
        current_content = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Try to detect headings (short lines that might be titles)
            # This is a heuristic - lines that are short and don't end with punctuation
            is_heading = (
                len(line) < 100 and
                not line.endswith('.') and
                not line.endswith(',') and
                not line.endswith(';') and
                (line[0].isupper() or line[0].isdigit() or line[0] in 'АБВГДЕЖЗИЙКЛМНОПРСТУФХЦЧШЩЬЫЪЭЮЯ')
            )
            
            # Check for numbered sections like "1.", "1.1", etc.
            import re
            if re.match(r'^\d+\.?\d*\.?\s', line) and len(line) < 100:
                is_heading = True
            
            if is_heading:
                # Save previous section if exists
                if current_section:
                    current_section["content"] = "\n".join(current_content)
                    structure["sections"].append(current_section)
                
                # Start new section
                current_section = {
                    "level": 1,
                    "title": line,
                    "content": ""
                }
                current_content = []
                
                # Set document title from first heading
                if not structure["title"]:
                    structure["title"] = line
                    logger.info(f"  - Document title: {line}")
            else:
                current_content.append(line)
        
        # Add last section
        if current_section:
            current_section["content"] = "\n".join(current_content)
            structure["sections"].append(current_section)
        
        logger.info(f"  - Found {len(structure['sections'])} sections")
        return structure
    
    def extract_tables(self) -> list:
        """Extract tables from PDF.
        
        Note: PDF table extraction is limited. For best results,
        use structured documents like Word or convert PDF to Word first.
        
        Returns:
            Empty list (PDF table extraction not implemented)
        """
        logger.warning("PDF table extraction is not implemented. Tables will be skipped.")
        return []
    
    def get_full_content(self) -> dict:
        """Get full PDF content including text and structure.
        
        Returns:
            Dictionary with all document content
        """
        logger.info("Getting full PDF content...")
        content = {
            "raw_text": self.extract_text(),
            "structure": self.extract_structure(),
            "tables": self.extract_tables()
        }
        logger.info("Full PDF content extraction completed")
        return content


class DocumentParser:
    """Parser for Microsoft Word documents (.doc, .docx)."""
    
    def __init__(self, file_path: str):
        """Initialize the document parser.
        
        Args:
            file_path: Path to the Word document
        """
        self.file_path = file_path
        self.document: Optional[Document] = None
        logger.debug(f"DocumentParser initialized for: {file_path}")
    
    def load_document(self) -> bool:
        """Load the Word document.
        
        Returns:
            True if document loaded successfully, False otherwise
        """
        logger.info(f"Loading document: {self.file_path}")
        try:
            self.document = Document(self.file_path)
            paragraphs_count = len(self.document.paragraphs)
            tables_count = len(self.document.tables)
            logger.info(f"Document loaded successfully:")
            logger.info(f"  - Paragraphs: {paragraphs_count}")
            logger.info(f"  - Tables: {tables_count}")
            return True
        except PackageNotFoundError:
            logger.error(f"Invalid Word document format: {self.file_path}")
            raise ValueError(f"Invalid Word document format: {self.file_path}")
        except Exception as e:
            logger.error(f"Error loading document: {str(e)}")
            raise RuntimeError(f"Error loading document: {str(e)}")
    
    def extract_text(self) -> str:
        """Extract all text from the document.
        
        Returns:
            Extracted text as a string
        """
        if not self.document:
            self.load_document()
        
        logger.info("Extracting text from document...")
        text_parts = []
        
        # Extract text from paragraphs
        for paragraph in self.document.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        paragraphs_with_text = len(text_parts)
        logger.info(f"  - Extracted text from {paragraphs_with_text} paragraphs")
        
        # Extract text from tables
        tables_text_count = 0
        for table in self.document.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))
                    tables_text_count += 1
        
        logger.info(f"  - Extracted text from {tables_text_count} table rows")
        
        result = "\n".join(text_parts)
        logger.info(f"Total text extracted: {len(result)} characters")
        return result
    
    def extract_structure(self) -> dict:
        """Extract document structure with headings and content.
        
        Returns:
            Dictionary with document structure
        """
        if not self.document:
            self.load_document()
        
        logger.info("Extracting document structure...")
        
        structure = {
            "title": "",
            "sections": []
        }
        
        current_section = None
        current_content = []
        
        for paragraph in self.document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            
            # Check if it's a heading
            if paragraph.style.name.startswith('Heading'):
                # Save previous section if exists
                if current_section:
                    current_section["content"] = "\n".join(current_content)
                    structure["sections"].append(current_section)
                
                # Start new section
                heading_level = int(paragraph.style.name.replace('Heading ', '')) if paragraph.style.name != 'Heading' else 1
                current_section = {
                    "level": heading_level,
                    "title": text,
                    "content": ""
                }
                current_content = []
                
                # Set document title from first heading
                if not structure["title"] and heading_level == 1:
                    structure["title"] = text
                    logger.info(f"  - Document title: {text}")
            else:
                # Add content to current section or as general content
                current_content.append(text)
        
        # Add last section
        if current_section:
            current_section["content"] = "\n".join(current_content)
            structure["sections"].append(current_section)
        
        logger.info(f"  - Found {len(structure['sections'])} sections")
        for section in structure['sections']:
            logger.info(f"    - Level {section['level']}: {section['title']}")
        
        return structure
    
    def extract_tables(self) -> list:
        """Extract all tables from the document.
        
        Returns:
            List of tables, each table is a list of rows
        """
        if not self.document:
            self.load_document()
        
        logger.info("Extracting tables from document...")
        
        tables = []
        for i, table in enumerate(self.document.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)
            logger.info(f"  - Table {i+1}: {len(table_data)} rows, {len(table_data[0]) if table_data else 0} columns")
        
        logger.info(f"Total tables extracted: {len(tables)}")
        return tables
    
    def get_full_content(self) -> dict:
        """Get full document content including text, structure, and tables.
        
        Returns:
            Dictionary with all document content
        """
        logger.info("Getting full document content...")
        content = {
            "raw_text": self.extract_text(),
            "structure": self.extract_structure(),
            "tables": self.extract_tables()
        }
        logger.info("Full content extraction completed")
        return content


def parse_document(file_path: str) -> dict:
    """Parse a document (Word or PDF) and return its content.
    
    Args:
        file_path: Path to the document (.doc, .docx, or .pdf)
        
    Returns:
        Dictionary with document content
    """
    logger.info(f"Starting document parsing: {file_path}")
    
    if not os.path.exists(file_path):
        logger.error(f"File not found: {file_path}")
        raise FileNotFoundError(f"File not found: {file_path}")
    
    file_size = os.path.getsize(file_path)
    logger.info(f"File size: {file_size} bytes")
    
    # Determine file type by extension
    _, ext = os.path.splitext(file_path.lower())
    
    if ext == '.pdf':
        logger.info("Detected PDF file, using PDF parser")
        parser = PDFParser(file_path)
    elif ext in ('.doc', '.docx'):
        logger.info(f"Detected Word file ({ext}), using Document parser")
        parser = DocumentParser(file_path)
    else:
        logger.error(f"Unsupported file type: {ext}")
        raise ValueError(f"Unsupported file type: {ext}. Supported types: .doc, .docx, .pdf")
    
    result = parser.get_full_content()
    
    logger.info("Document parsing completed successfully")
    return result
